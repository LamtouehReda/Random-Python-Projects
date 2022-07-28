import docx
from docx.shared import Inches
from tkinter import filedialog
import tkinter as tk
# import aspose.words as aw
import docx2pdf 


class Window:
	def __init__(self,root,title,geometry):
		self.root=root
		self.root.title(title)
		self.root.geometry(geometry)

		self.label=tk.Label(self.root,text='Auto Math Homework Generator By Reda Lamtoueh')
		self.label.grid(row=0,column=0)

		self.paths=[]

		self.BrowseBtn=tk.Button(self.root,text='Select Images',command=self.Browse )
		self.BrowseBtn.grid(row=1,column=0)

		self.frame1=tk.Frame(self.root)

		self.label1=tk.Label(self.frame1,text='Title : ')
		self.label1.grid(row=0,column=0)

		self.idNameEntry=tk.Entry(self.frame1)
		self.idNameEntry.grid(row=0,column=1)

		self.frame1.grid(row=2,column=0)

		self.submitBtn=tk.Button(self.root,text='Generate',command=self.generateDoc)
		self.submitBtn.grid(row=3,column=0)

		self.root.mainloop()

	def Browse(self):
		images=filedialog.askopenfilenames()
		for img in images:
			self.paths.append(img)




	def generateDoc(self):
		title=self.idNameEntry.get()
		doc=docx.Document()
		doc.add_heading(title)

		for path in self.paths:
			doc.add_picture(path,width=Inches(6))

		doc.save(f'{title}.docx')

		#save the doc as pdf as well
		# doc=aw.Document(f'{title}.docx')
		# doc.save(f'{title}.pdf')
		docx2pdf.convert(f'{title}.docx',f'{title}.pdf')
		docx2pdf.convert('')


root=tk.Tk()

window=Window(root,'Homework Generator','300x200')